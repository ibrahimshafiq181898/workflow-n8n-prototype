from flask import Flask, jsonify, request,send_file, abort
from flask_cors import CORS
import requests
from io import BytesIO
from docx import Document
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

N8N_WEBHOOK_URL = "https://tjans.app.n8n.cloud/webhook/file-processing"

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    # Send file to n8n workflow
    files = {'file': (file.filename, file.stream, file.content_type)}
    response = requests.post(N8N_WEBHOOK_URL, files=files)

    # Return the response from n8n
    if response.ok:
        return jsonify(response.json())
    else:
        return jsonify({"error": "n8n workflow failed"}), 500
@app.route("/generate-doc", methods=["POST"])
def generate_doc():
    data = request.json
    if not data or "texts" not in data:
        abort(400, 'Expecting JSON with key "texts" (list of strings)')

    texts = data["texts"]
    if not isinstance(texts, list):
        abort(400, '"texts" must be a list of strings')

    doc = Document()

    for t in texts:
        # Split by lines in case of multi-line text
        lines = t.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph("")  # Blank line
                continue

            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line[0:3].isdigit() and line[3:5] == ". ":
                doc.add_paragraph(line[5:], style="List Number")
            else:
                doc.add_paragraph(line)  # Normal paragraph

        # Add extra blank line between proposals
        doc.add_paragraph("")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name="proposals.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    data = request.json
    if not data or "texts" not in data:
        abort(400, 'Expecting JSON with key "texts" (list of strings)')

    texts = data["texts"]
    if not isinstance(texts, list):
        abort(400, '"texts" must be a list of strings')

    doc = Document()
    for t in texts:
        doc.add_paragraph(t)       # Add the proposal
        doc.add_paragraph("")      # Blank line between proposals

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name="proposals.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
if __name__ == "__main__":
    app.run(port=5000, debug=True)
